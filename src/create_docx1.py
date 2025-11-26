import json
import os
import PIL.Image

if not hasattr(PIL.Image, 'ANTIALIAS'):
    PIL.Image.ANTIALIAS = PIL.Image.LANCZOS

import platform
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt

# PDF helpers
from fpdf import FPDF  # fpdf2 (pure Python)

# Optional: docx2pdf (Windows/Mac) — used if available
try:
    from docx2pdf import convert as docx2pdf_convert
    _HAS_DOCX2PDF = True
except Exception:
    _HAS_DOCX2PDF = False


def clean_text(text):
    """
    Sanitizes text for the FPDF fallback. 
    Replaces special Unicode characters that might cause crashes in Latin-1 fonts
    if a custom Unicode font is not found.
    """
    if not isinstance(text, str):
        return str(text) if text is not None else ""
    
    replacements = {
        '\u2013': '-',   # En-dash
        '\u2014': '--',  # Em-dash
        '\u2018': "'",   # Left single quote
        '\u2019': "'",   # Right single quote
        '\u201c': '"',   # Left double quote
        '\u201d': '"',   # Right double quote
        '\u2022': '*',   # Bullet
        '\u00a0': ' '    # Non-breaking space
    }
    
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    
    return text


def fill_template(template_path, output_docx_path, output_pdf_path, resume_json):
    """
    Fill the DOCX template and generate a PDF.
    """
    print("▶ fill_template start")

    # --- 1. Parse Data Robustly ---
    if isinstance(resume_json, str):
        try:
            data = json.loads(resume_json)
        except json.JSONDecodeError:
            print("❌ Error: Invalid JSON string provided.")
            return
    else:
        data = resume_json if resume_json else {}

    doc = Document(template_path)

    # --- Helper Inner Functions ---
    def find_heading_idx(heading_text):
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip().lower() == heading_text.lower():
                return idx
        return None

    def clear_paragraph(para: Paragraph):
        for run in para.runs:
            run.text = ""
        para.text = ""

    def add_after_heading(heading, lines, bullet=False):
        idx = find_heading_idx(heading)
        if idx is not None:
            # skip placeholder empty paragraphs immediately after header
            while idx + 1 < len(doc.paragraphs) and not doc.paragraphs[idx + 1].text.strip():
                idx += 1
            for line in lines:
                new_para = doc.add_paragraph(line, style="List Bullet" if bullet else None)
                doc.paragraphs[idx]._element.addnext(new_para._element)
                idx += 1

    # --- 2. Fill Core Fields ---

    # Full Name
    for para in doc.paragraphs:
        if "full name" in para.text.lower():
            full_name = f"{data.get('first_name', '')} {data.get('last_name', '')}".strip()
            clear_paragraph(para)
            para.add_run(full_name)
            break

    # Career Summary
    summary_text = data.get("career_summary", "")
    if summary_text:
        # Split by periods to create a list, or just insert as block
        # Using simple block insertion here to maintain flow
        idx = find_heading_idx("Career Summary")
        if idx is not None:
            if idx + 1 < len(doc.paragraphs):
                clear_paragraph(doc.paragraphs[idx + 1])
                doc.paragraphs[idx + 1].add_run(summary_text)

    # Expertise (Updated: Joined by Pipe |)
    idx = find_heading_idx("Expertise")
    if idx is not None and idx + 1 < len(doc.paragraphs):
        clear_paragraph(doc.paragraphs[idx + 1])
        expertise_list = data.get("expertise", [])
        if expertise_list:
            doc.paragraphs[idx + 1].add_run(" | ".join(expertise_list))

    # Technical Skills (Updated: Joined by Pipe |)
    idx = find_heading_idx("Technical Skills")
    if idx is not None and idx + 1 < len(doc.paragraphs):
        clear_paragraph(doc.paragraphs[idx + 1])
        skills_list = data.get("technical_skills", [])
        if skills_list:
            doc.paragraphs[idx + 1].add_run(" | ".join(skills_list))

    # Professional Experience
    idx = find_heading_idx("Professional Experience")
    if idx is not None:
        # We insert in reverse order so they appear top-down correctly when using addnext
        # Or standard order if we increment idx. The original code incremented idx, 
        # so we iterate normally.
        for job in data.get("professional_experience", []):
            if not isinstance(job, dict): continue

            # Header: Title | Company | Location
            header_parts = [job.get("title", ""), job.get("company", ""), job.get("location", "")]
            header_parts = [p for p in header_parts if p]
            
            if header_parts:
                hp = doc.add_paragraph(" | ".join(header_parts))
                if hp.runs: hp.runs[0].bold = True
                doc.paragraphs[idx]._element.addnext(hp._element)
                idx += 1

            # Dates
            date_parts = [job.get("start_date", ""), job.get("end_date", "")]
            date_parts = [d for d in date_parts if d]
            if date_parts:
                dp = doc.add_paragraph(" – ".join(date_parts)) # Unicode En-dash
                if dp.runs: dp.runs[0].italic = True
                doc.paragraphs[idx]._element.addnext(dp._element)
                idx += 1

            # Description
            if job.get("description"):
                dp2 = doc.add_paragraph(job["description"])
                doc.paragraphs[idx]._element.addnext(dp2._element)
                idx += 1

            # Achievements
            achievements = job.get("achievements", [])
            if isinstance(achievements, list):
                for ach in achievements:
                    if ach.strip():
                        a = doc.add_paragraph(ach, style="List Bullet")
                        doc.paragraphs[idx]._element.addnext(a._element)
                        idx += 1
            
            # Spacer
            spacer = doc.add_paragraph("")
            doc.paragraphs[idx]._element.addnext(spacer._element)
            idx += 1

    # Education
    edu_lines = []
    for edu in data.get("education", []):
        parts = []
        if edu.get("degree"): parts.append(edu["degree"])
        if edu.get("field"): parts.append(f"in {edu['field']}")
        if edu.get("institution"): parts.append(f"| {edu['institution']}")
        if edu.get("end_year"): parts.append(f"| {edu['end_year']}")
        
        line = " ".join(parts)
        if line:
            edu_lines.append(line)
    
    if edu_lines:
        add_after_heading("Education", edu_lines, bullet=True)

    # --- 3. Handle Certifications (Remove if empty) ---
    certifications = data.get("certifications", [])
    cert_lines = [c.strip() for c in certifications if c and c.strip()]

    if not cert_lines:
        # Remove the section entirely if empty
        indices_to_remove = []
        for i, p in enumerate(doc.paragraphs):
            if "certification" in p.text.strip().lower():
                indices_to_remove.append(i)
                # Check for subsequent empty placeholder
                if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i+1].text.strip():
                    indices_to_remove.append(i+1)
                break # Assuming only one Cert section
        
        for idx in sorted(indices_to_remove, reverse=True):
            p = doc.paragraphs[idx]._element
            p.getparent().remove(p)
    else:
        # Fill it
        add_after_heading("CERTIFICATION", cert_lines, bullet=True)
    
    # --- 4. Save DOCX ---
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    doc.save(output_docx_path)
    print(f"✅ DOCX saved: {output_docx_path}")

    # --- 5. Generate PDF ---
    create_pdf_from_docx(output_docx_path, output_pdf_path)


def find_dejavu_ttf():
    """Return a path to a DejaVuSans.ttf if present on the system or None."""
    candidates = [
        # Linux common
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        # macOS common
        "/Library/Fonts/DejaVuSans.ttf",
        # Windows common locations (may vary)
        "C:\\Windows\\Fonts\\DejaVuSans.ttf",
        # Local project fallback
        os.path.join(os.getcwd(), "DejaVuSans.ttf"),
    ]
    for p in candidates:
        if p and os.path.exists(p):
            return p
    return None


def create_pdf_from_docx(docx_path: str, output_pdf_path: str):
    """
    Create a PDF from a DOCX file.
    Priority 1: docx2pdf (High Fidelity, preserves DOCX layout)
    Priority 2: fpdf (Low Fidelity, text-dump fallback)
    """
    
    # 1. Try docx2pdf
    if _HAS_DOCX2PDF and platform.system() in ("Windows", "Darwin"):
        try:
            print("ℹ Using docx2pdf for conversion (best fidelity)")
            os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
            docx2pdf_convert(docx_path, output_pdf_path)
            return
        except Exception as e:
            print(f"⚠ docx2pdf conversion failed: {e} — falling back to text-render PDF")

    # 2. Fallback: fpdf text-render
    print("ℹ Falling back to fpdf text-render conversion")
    doc = Document(docx_path)

    pdf = FPDF(format='A4', unit='mm')
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Font handling
    dejavu = find_dejavu_ttf()
    use_unicode_font = False
    
    if dejavu:
        try:
            # Note: fpdf version < 2.5 uses 'uni=True', newer versions behave differently.
            # Assuming standard fpdf usage here.
            pdf.add_font("DejaVu", "", dejavu, uni=True)
            pdf.set_font("DejaVu", size=11)
            use_unicode_font = True
            print(f"ℹ Registered DejaVu font at: {dejavu}")
        except Exception as e:
            print(f"⚠ Failed to register DejaVu font: {e} — using built-in font")
            pdf.set_font("Arial", size=11)
    else:
        print("⚠ DejaVuSans.ttf not found — unicode characters may not render correctly")
        pdf.set_font("Arial", size=11)

    # Render Text
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            pdf.ln(4)
            continue
        
        # If we aren't using a custom Unicode font, we MUST clean the text 
        # to remove chars like En-dash (–) which will crash standard Arial
        if not use_unicode_font:
            text = clean_text(text)
            
        try:
            pdf.multi_cell(0, 6, txt=text)
            pdf.ln(1)
        except Exception as e:
            # Final safety net: try one more time completely stripping non-ascii
            print(f"⚠ Error printing line: {e}. Attempting ASCII fallback.")
            safe_text = text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, txt=safe_text)

    os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
    pdf.output(output_pdf_path)
    if os.path.exists(output_pdf_path):
        print(f"✅ PDF saved: {output_pdf_path}")


# Execution Entry Point
if __name__ == '__main__':
    # Update these paths for your local testing
    tpl = 'data/template/Resume_Template_PGI.docx'
    out_docx = 'data/output/doc/test.docx'
    out_pdf = 'data/output/pdf/test.pdf'
    
    # You can pass JSON string or dict here
    sample_json = {} 
    
    if os.path.exists(tpl):
        fill_template(tpl, out_docx, out_pdf, sample_json)
    else:
        print(f'⚠ Template not found at {tpl} — update the path for manual testing')