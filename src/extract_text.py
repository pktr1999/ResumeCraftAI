# from PyPDF2 import PdfReader

# def extract_text_from_pdf(pdf_path):
#     reader = PdfReader(pdf_path)
#     text=""
#     for page in reader.pages:
#         text += page.extract_text() + "\n"
#     return text


# from pdf2image import convert_from_path, convert_from_bytes
# import pytesseract
# from PyPDF2 import PdfReader
# from docx import Document
# from io import BytesIO
# import os

# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# # ----------- OCR for scanned PDF (Poppler + pytesseract) -----------
# def extract_text_from_pdf_ocr(pdf_source):
#     print("Running OCR with pdf2image + pytesseract...")

#     # Case: Streamlit BytesIO upload
#     if isinstance(pdf_source, BytesIO):
#         print("Detected BytesIO — using convert_from_bytes")
#         pdf_bytes = pdf_source.getvalue()
#         images = convert_from_bytes(pdf_bytes, dpi=300)

#     else:
#         print("Detected file path — using convert_from_path")
#         images = convert_from_path(str(pdf_source), dpi=300)

#     full_text = ""
#     for img in images:
#         text = pytesseract.image_to_string(img)
#         full_text += text + "\n"

#     return full_text


# # ----------- Normal PDF text extraction -----------
# def extract_text_from_pdf(pdf_source):
#     print("Extracting with PyPDF2...")

#     try:
#         if isinstance(pdf_source, BytesIO):
#             reader = PdfReader(pdf_source)
#         else:
#             reader = PdfReader(str(pdf_source))
#     except Exception as e:
#         print(f"PyPDF2 failed ({e}) — switching to OCR.")
#         return extract_text_from_pdf_ocr(pdf_source)

#     text = ""
#     for page in reader.pages:
#         extracted = page.extract_text()
#         if extracted:
#             text += extracted + "\n"

#     if text.strip() == "":
#         print("PDF contains no text — running OCR.")
#         return extract_text_from_pdf_ocr(pdf_source)

#     return text


# # ----------- DOCX extraction -----------
# def extract_text_from_docx(docx_source):
#     print("Extracting DOCX...")
#     doc = Document(docx_source)
#     return "\n".join([p.text for p in doc.paragraphs])


# # ----------- Unified function -----------
# def extract_text(file_input):
#     print("extract_text called...")

#     # Case: Streamlit BytesIO upload
#     if isinstance(file_input, BytesIO):
#         print("Detected BytesIO upload")

#         file_input.seek(0)
#         try:
#             return extract_text_from_pdf(file_input)
#         except Exception as e:
#             print(f"PDF extraction failed ({e}), trying DOCX...")
#             file_input.seek(0)
#             return extract_text_from_docx(file_input)

#     # Case: File path (string, bytes, Path object)
#     elif isinstance(file_input, (str, bytes, os.PathLike)):
#         ext = os.path.splitext(str(file_input))[1].lower()

#         if ext == ".pdf":
#             return extract_text_from_pdf(file_input)
#         elif ext == ".docx":
#             return extract_text_from_docx(file_input)
#         else:
#             raise ValueError(f"Unsupported file extension: {ext}")

#     else:
#         raise TypeError(f"Unsupported input type: {type(file_input)}")

import os
from io import BytesIO
import fitz  # PyMuPDF
import easyocr
from PyPDF2 import PdfReader
from docx import Document

# Initialize EasyOCR reader ONCE (expensive)
reader = easyocr.Reader(['en'], gpu=False)


# ======================================================
# 1) OCR for scanned PDFs (PyMuPDF + EasyOCR)
# ======================================================
def extract_text_from_pdf_ocr(pdf_source):
    print("Running OCR with PyMuPDF + EasyOCR...")

    # Handle BytesIO vs file path:
    if isinstance(pdf_source, BytesIO):
        pdf_bytes = pdf_source.getvalue()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    else:
        doc = fitz.open(pdf_source)

    text = ""

    for page in doc:
        # Render page → image
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")

        # EasyOCR on the image bytes
        result = reader.readtext(img_bytes, detail=0, paragraph=True)

        page_text = "\n".join(result)
        text += page_text + "\n"

    doc.close()
    return text


# ======================================================
# 2) Normal PDF text extraction (PyPDF2)
# ======================================================
def extract_text_from_pdf(pdf_source):
    print("Extracting text using PyPDF2...")

    try:
        if isinstance(pdf_source, BytesIO):
            reader = PdfReader(pdf_source)
        else:
            reader = PdfReader(str(pdf_source))
    except Exception as e:
        print(f"PyPDF2 failed ({e}) → Using OCR instead.")
        return extract_text_from_pdf_ocr(pdf_source)

    full_text = ""

    for page in reader.pages:
        txt = page.extract_text()
        if txt:
            full_text += txt + "\n"

    # If no selectable text → scanned PDF → run OCR
    if full_text.strip() == "":
        print("PDF has no extractable text → Running OCR...")
        return extract_text_from_pdf_ocr(pdf_source)

    return full_text


# ======================================================
# 3) DOCX extraction
# ======================================================
def extract_text_from_docx(docx_source):
    print("Extracting text from DOCX...")
    doc = Document(docx_source)
    return "\n".join([p.text for p in doc.paragraphs])


# ======================================================
# 4) Main dispatcher (BytesIO or file path)
# ======================================================
def extract_text(file_input):
    print("extract_text() called...")

    # Case 1: Streamlit file upload (BytesIO)
    if isinstance(file_input, BytesIO):
        print("Detected BytesIO upload...")
        file_input.seek(0)

        try:
            return extract_text_from_pdf(file_input)
        except Exception as e:
            print(f"PDF parse failed ({e}) → Trying DOCX instead...")
            file_input.seek(0)
            return extract_text_from_docx(file_input)

    # Case 2: Local file path
    elif isinstance(file_input, (str, bytes, os.PathLike)):
        ext = os.path.splitext(str(file_input))[1].lower()

        if ext == ".pdf":
            return extract_text_from_pdf(file_input)

        elif ext == ".docx":
            return extract_text_from_docx(file_input)

        else:
            raise ValueError(f"Unsupported file type: {ext}")

    # Unsupported input type
    else:
        raise TypeError(f"Unsupported input type: {type(file_input)}")
