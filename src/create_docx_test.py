# # import json
# # import os
# # import platform
# # from docx import Document
# # from docx.text.paragraph import Paragraph
# # from docx.shared import Pt

# # # PDF helpers
# # from fpdf import FPDF

# # # Optional: docx2pdf (Windows/Mac) — used if available
# # try:
# #     from docx2pdf import convert as docx2pdf_convert
# #     _HAS_DOCX2PDF = True
# # except Exception:
# #     _HAS_DOCX2PDF = False


# # def clean_text(text):
# #     """
# #     Sanitizes text for FPDF (Latin-1) fallback. 
# #     Replaces special Unicode characters that might cause crashes.
# #     """
# #     if not isinstance(text, str):
# #         return str(text) if text is not None else ""
    
# #     replacements = {
# #         '\u2013': '-',   # En-dash
# #         '\u2014': '--',  # Em-dash
# #         '\u2018': "'",   # Left single quote
# #         '\u2019': "'",   # Right single quote
# #         '\u201c': '"',   # Left double quote
# #         '\u201d': '"',   # Right double quote
# #         '\u2022': '*',   # Bullet
# #         '\u00a0': ' '    # Non-breaking space
# #     }
    
# #     for char, replacement in replacements.items():
# #         text = text.replace(char, replacement)
    
# #     return text


# # def fill_template(template_path, output_docx_path, output_pdf_path, resume_json):
# #     """
# #     Fill the DOCX template (python-docx) and generate a readable PDF.
# #     """
# #     print("▶ fill_template start")

# #     # --- 1. Parse Data Robustly ---
# #     if isinstance(resume_json, str):
# #         try:
# #             data = json.loads(resume_json)
# #         except json.JSONDecodeError:
# #             print("❌ Error: Invalid JSON string provided.")
# #             return
# #     else:
# #         # Handle dict or None
# #         data = resume_json if resume_json else {}

# #     try:
# #         doc = Document(template_path)
# #     except Exception as e:
# #         print(f"❌ Error loading template: {e}")
# #         return

# #     # --- Helper Inner Functions ---
# #     def find_heading_idx(heading_text):
# #         for idx, para in enumerate(doc.paragraphs):
# #             if para.text.strip().lower() == heading_text.lower():
# #                 return idx
# #         return None

# #     def clear_paragraph(para: Paragraph):
# #         for run in para.runs:
# #             run.text = ""
# #         para.text = ""

# #     def add_after_heading(heading, lines, bullet=False):
# #         idx = find_heading_idx(heading)
# #         if idx is not None:
# #             # Skip placeholder empty paragraphs immediately after header
# #             while idx + 1 < len(doc.paragraphs) and not doc.paragraphs[idx + 1].text.strip():
# #                 idx += 1
# #             for line in lines:
# #                 if not line.strip(): continue
# #                 new_para = doc.add_paragraph(line, style="List Bullet" if bullet else None)
# #                 # Insert paragraph after the current index
# #                 doc.paragraphs[idx]._element.addnext(new_para._element)
# #                 idx += 1

# #     # --- 2. Fill Core Fields ---

# #     # Full Name
# #     for para in doc.paragraphs:
# #         if "full name" in para.text.lower():
# #             full_name = f"{data.get('first_name', '')} {data.get('last_name', '')}".strip()
# #             clear_paragraph(para)
# #             para.add_run(full_name)
# #             break

# #     # Career Summary
# #     summary_text = data.get("career_summary", "")
# #     if summary_text:
# #         idx = find_heading_idx("Career Summary")
# #         if idx is not None and idx + 1 < len(doc.paragraphs):
# #             clear_paragraph(doc.paragraphs[idx + 1])
# #             doc.paragraphs[idx + 1].add_run(summary_text)

# #     # Expertise (Joined by Pipe |)
# #     idx = find_heading_idx("Expertise")
# #     if idx is not None and idx + 1 < len(doc.paragraphs):
# #         clear_paragraph(doc.paragraphs[idx + 1])
# #         expertise_list = data.get("expertise", [])
# #         if expertise_list:
# #             doc.paragraphs[idx + 1].add_run(" | ".join(expertise_list))

# #     # Technical Skills (Joined by Pipe |)
# #     idx = find_heading_idx("Technical Skills")
# #     if idx is not None and idx + 1 < len(doc.paragraphs):
# #         clear_paragraph(doc.paragraphs[idx + 1])
# #         skills_list = data.get("technical_skills", [])
# #         if skills_list:
# #             doc.paragraphs[idx + 1].add_run(" | ".join(skills_list))

# #     # Professional Experience
# #     idx = find_heading_idx("Professional Experience")
# #     if idx is not None:
# #         # We iterate through the list. Note: insertion logic (addnext) inserts *after* the header.
# #         # To keep chronological order correct relative to the header, we can just insert sequentially
# #         # pushing previous content down, OR simpler: let's insert them one by one.
# #         # Since 'addnext' pushes existing content down, we must keep incrementing idx to append to the bottom of the section.
        
# #         for job in data.get("professional_experience", []):
# #             if not isinstance(job, dict): continue

# #             # Header: Title | Company | Location
# #             header_parts = [job.get("title", ""), job.get("company", ""), job.get("location", "")]
# #             header_parts = [p for p in header_parts if p]
            
# #             if header_parts:
# #                 hp = doc.add_paragraph(" | ".join(header_parts))
# #                 if hp.runs: hp.runs[0].bold = True
# #                 doc.paragraphs[idx]._element.addnext(hp._element)
# #                 idx += 1

# #             # Dates
# #             date_parts = [job.get("start_date", ""), job.get("end_date", "")]
# #             date_parts = [d for d in date_parts if d]
# #             if date_parts:
# #                 dp = doc.add_paragraph(" – ".join(date_parts)) # Unicode En-dash
# #                 if dp.runs: dp.runs[0].italic = True
# #                 doc.paragraphs[idx]._element.addnext(dp._element)
# #                 idx += 1

# #             # Description
# #             if job.get("description"):
# #                 dp2 = doc.add_paragraph(job["description"])
# #                 doc.paragraphs[idx]._element.addnext(dp2._element)
# #                 idx += 1

# #             # Achievements
# #             achievements = job.get("achievements", [])
# #             if isinstance(achievements, list):
# #                 for ach in achievements:
# #                     if ach.strip():
# #                         a = doc.add_paragraph(ach, style="List Bullet")
# #                         doc.paragraphs[idx]._element.addnext(a._element)
# #                         idx += 1
            
# #             # Spacer
# #             spacer = doc.add_paragraph("")
# #             doc.paragraphs[idx]._element.addnext(spacer._element)
# #             idx += 1

# #     # Education
# #     edu_lines = []
# #     for edu in data.get("education", []):
# #         parts = []
# #         if edu.get("degree"): parts.append(edu["degree"])
# #         if edu.get("field"): parts.append(f"in {edu['field']}")
# #         if edu.get("institution"): parts.append(f"| {edu['institution']}")
# #         if edu.get("end_year"): parts.append(f"| {edu['end_year']}")
        
# #         line = " ".join(parts)
# #         if line:
# #             edu_lines.append(line)
    
# #     if edu_lines:
# #         add_after_heading("Education", edu_lines, bullet=True)

# #     # --- 3. Handle Certifications (Remove if empty) ---
# #     certifications = data.get("certifications", [])
# #     cert_lines = [c.strip() for c in certifications if c and c.strip()]

# #     if not cert_lines:
# #         # Remove the section entirely if empty
# #         indices_to_remove = []
# #         for i, p in enumerate(doc.paragraphs):
# #             if "certification" in p.text.strip().lower():
# #                 indices_to_remove.append(i)
# #                 # Check for subsequent empty placeholder
# #                 if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i+1].text.strip():
# #                     indices_to_remove.append(i+1)
# #                 break 
        
# #         for idx in sorted(indices_to_remove, reverse=True):
# #             p = doc.paragraphs[idx]._element
# #             p.getparent().remove(p)
# #     else:
# #         # Fill it
# #         add_after_heading("CERTIFICATION", cert_lines, bullet=True)
    
# #     # --- 4. Save DOCX ---
# #     os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
# #     doc.save(output_docx_path)
# #     print(f"✅ DOCX saved: {output_docx_path}")

# #     # --- 5. Generate PDF ---
# #     create_pdf_from_docx(output_docx_path, output_pdf_path)


# # def find_dejavu_ttf():
# #     """Return a path to a DejaVuSans.ttf if present on the system or None."""
# #     candidates = [
# #         "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
# #         "/Library/Fonts/DejaVuSans.ttf",
# #         "C:\\Windows\\Fonts\\DejaVuSans.ttf",
# #         os.path.join(os.getcwd(), "DejaVuSans.ttf"),
# #     ]
# #     for p in candidates:
# #         if p and os.path.exists(p):
# #             return p
# #     return None


# # def create_pdf_from_docx(docx_path: str, output_pdf_path: str):
# #     """
# #     Create a PDF from a DOCX file.
# #     Priority 1: docx2pdf (High Fidelity, preserves DOCX layout)
# #     Priority 2: fpdf (Low Fidelity, text-dump fallback)
# #     """
    
# #     # 1. Try docx2pdf
# #     if _HAS_DOCX2PDF and platform.system() in ("Windows", "Darwin"):
# #         try:
# #             print("ℹ Using docx2pdf for conversion (best fidelity)")
# #             os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
# #             docx2pdf_convert(docx_path, output_pdf_path)
# #             return
# #         except Exception as e:
# #             print(f"⚠ docx2pdf conversion failed: {e} — falling back to text-render PDF")

# #     # 2. Fallback: fpdf text-render
# #     print("ℹ Falling back to fpdf text-render conversion")
    
# #     try:
# #         doc = Document(docx_path)
# #     except Exception as e:
# #         print(f"❌ Failed to read temporary DOCX for PDF generation: {e}")
# #         return

# #     pdf = FPDF(format='A4', unit='mm')
# #     pdf.set_auto_page_break(auto=True, margin=15)
# #     pdf.add_page()

# #     # Font handling
# #     dejavu = find_dejavu_ttf()
# #     use_unicode_font = False
    
# #     if dejavu:
# #         try:
# #             # Note: fpdf version compatibility check
# #             # Some versions use 'uni=True', others detect automatically or don't support it
# #             # We try standard add_font
# #             pdf.add_font("DejaVu", "", dejavu, uni=True)
# #             pdf.set_font("DejaVu", size=11)
# #             use_unicode_font = True
# #             print(f"ℹ Registered DejaVu font at: {dejavu}")
# #         except Exception:
# #             try:
# #                 # Fallback for different fpdf versions (remove uni=True)
# #                 pdf.add_font("DejaVu", "", dejavu)
# #                 pdf.set_font("DejaVu", size=11)
# #                 use_unicode_font = True
# #             except Exception as e:
# #                 print(f"⚠ Failed to register DejaVu font: {e} — using built-in Arial")
# #                 pdf.set_font("Arial", size=11)
# #     else:
# #         print("⚠ DejaVuSans.ttf not found — unicode characters may not render correctly")
# #         pdf.set_font("Arial", size=11)

# #     # --- ERROR FIX: Explicitly Calculate Width ---
# #     # fpdf2 uses .epw (Effective Page Width), older fpdf uses calculations.
# #     try:
# #         if hasattr(pdf, 'epw'):
# #             effective_width = pdf.epw
# #         else:
# #             effective_width = pdf.w - pdf.l_margin - pdf.r_margin
# #     except Exception:
# #         # Fallback to standard A4 width minus standard margins (210 - 20)
# #         effective_width = 190

# #     # Render Text
# #     for para in doc.paragraphs:
# #         text = para.text.strip()
# #         if not text:
# #             pdf.ln(4)
# #             continue
        
# #         # If no Unicode font, clean text to avoid crashes
# #         if not use_unicode_font:
# #             text = clean_text(text)
            
# #         try:
# #             # Use explicit width to prevent "Not enough horizontal space" error
# #             pdf.multi_cell(w=effective_width, h=6, txt=text)
# #             pdf.ln(1)
# #         except Exception as e:
# #             print(f"⚠ Error printing line: {e}. Attempting simplified fallback.")
# #             try:
# #                 # Absolute panic fallback: ASCII only, small width
# #                 safe_text = text.encode('ascii', 'ignore').decode('ascii')
# #                 pdf.multi_cell(w=180, h=6, txt=safe_text)
# #                 pdf.ln(1)
# #             except Exception:
# #                 pass # Skip line if it still fails

# #     os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
# #     pdf.output(output_pdf_path)
# #     if os.path.exists(output_pdf_path):
# #         print(f"✅ PDF saved: {output_pdf_path}")


# # import json
# # import os
# # import platform
# # import re
# # from docx import Document
# # from docx.text.paragraph import Paragraph
# # from docx.shared import Pt
# # from docx.enum.text import WD_ALIGN_PARAGRAPH

# # # PDF helpers
# # from fpdf import FPDF

# # # Optional: docx2pdf (Windows/Mac) — used if available
# # try:
# #     from docx2pdf import convert as docx2pdf_convert
# #     _HAS_DOCX2PDF = True
# # except Exception:
# #     _HAS_DOCX2PDF = False


# # def clean_text(text):
# #     """
# #     Sanitizes text for FPDF (Latin-1) fallback.
# #     Replaces special Unicode characters that might cause crashes.
# #     """
# #     if not isinstance(text, str):
# #         return str(text) if text is not None else ""
    
# #     replacements = {
# #         '\u2013': '-',   # En-dash
# #         '\u2014': '--',  # Em-dash
# #         '\u2018': "'",   # Left single quote
# #         '\u2019': "'",   # Right single quote
# #         '\u201c': '"',   # Left double quote
# #         '\u201d': '"',   # Right double quote
# #         '\u2022': '*',   # Bullet
# #         '\u00a0': ' '    # Non-breaking space
# #     }
    
# #     for char, replacement in replacements.items():
# #         text = text.replace(char, replacement)
    
# #     return text


# # def fill_template(template_path, output_docx_path, output_pdf_path, resume_json):
# #     """
# #     Fill the DOCX template (python-docx) and generate a readable PDF
# #     that mimics the DOCX formatting.
# #     """
# #     print("▶ fill_template start")

# #     # --- 1. Parse Data Robustly ---
# #     if isinstance(resume_json, str):
# #         try:
# #             data = json.loads(resume_json)
# #         except json.JSONDecodeError:
# #             print("❌ Error: Invalid JSON string provided.")
# #             return
# #     else:
# #         data = resume_json if resume_json else {}

# #     try:
# #         doc = Document(template_path)
# #     except Exception as e:
# #         print(f"❌ Error loading template: {e}")
# #         return

# #     # --- Helper Inner Functions ---
# #     def find_heading_idx(heading_text):
# #         for idx, para in enumerate(doc.paragraphs):
# #             if para.text.strip().lower() == heading_text.lower():
# #                 return idx
# #         return None

# #     def clear_paragraph(para: Paragraph):
# #         for run in para.runs:
# #             run.text = ""
# #         para.text = ""

# #     def add_after_heading(heading, lines, bullet=False):
# #         idx = find_heading_idx(heading)
# #         if idx is not None:
# #             # Skip placeholder empty paragraphs immediately after header
# #             while idx + 1 < len(doc.paragraphs) and not doc.paragraphs[idx + 1].text.strip():
# #                 idx += 1
# #             for line in lines:
# #                 if not line.strip(): continue
# #                 # We apply "List Bullet" style here so the PDF generator can detect it later
# #                 style = "List Bullet" if bullet else None
# #                 new_para = doc.add_paragraph(line, style=style)
# #                 doc.paragraphs[idx]._element.addnext(new_para._element)
# #                 idx += 1

# #     # --- 2. Fill Core Fields ---

# #     # Full Name
# #     for para in doc.paragraphs:
# #         if "full name" in para.text.lower():
# #             full_name = f"{data.get('first_name', '')} {data.get('last_name', '')}".strip()
# #             clear_paragraph(para)
# #             para.add_run(full_name)
# #             # Explicitly center and bold the name for DOCX
# #             para.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #             if para.runs: para.runs[0].bold = True
# #             break

# #     # Career Summary
# #     summary_text = data.get("career_summary", "")
# #     if summary_text:
# #         idx = find_heading_idx("Career Summary")
# #         if idx is not None and idx + 1 < len(doc.paragraphs):
# #             clear_paragraph(doc.paragraphs[idx + 1])
# #             doc.paragraphs[idx + 1].add_run(summary_text)

# #     # Expertise (Joined by Pipe |)
# #     idx = find_heading_idx("Expertise")
# #     if idx is not None and idx + 1 < len(doc.paragraphs):
# #         clear_paragraph(doc.paragraphs[idx + 1])
# #         expertise_list = data.get("expertise", [])
# #         if expertise_list:
# #             doc.paragraphs[idx + 1].add_run(" | ".join(expertise_list))

# #     # Technical Skills (Joined by Pipe |)
# #     idx = find_heading_idx("Technical Skills")
# #     if idx is not None and idx + 1 < len(doc.paragraphs):
# #         clear_paragraph(doc.paragraphs[idx + 1])
# #         skills_list = data.get("technical_skills", [])
# #         if skills_list:
# #             doc.paragraphs[idx + 1].add_run(" | ".join(skills_list))

# #     # Professional Experience
# #     idx = find_heading_idx("Professional Experience")
# #     if idx is not None:
# #         for job in data.get("professional_experience", []):
# #             if not isinstance(job, dict): continue

# #             # Header: Title | Company | Location
# #             header_parts = [job.get("title", ""), job.get("company", ""), job.get("location", "")]
# #             header_parts = [p for p in header_parts if p]
            
# #             if header_parts:
# #                 hp = doc.add_paragraph(" | ".join(header_parts))
# #                 if hp.runs: hp.runs[0].bold = True
# #                 doc.paragraphs[idx]._element.addnext(hp._element)
# #                 idx += 1

# #             # Dates
# #             date_parts = [job.get("start_date", ""), job.get("end_date", "")]
# #             date_parts = [d for d in date_parts if d]
# #             if date_parts:
# #                 dp = doc.add_paragraph(" – ".join(date_parts)) 
# #                 if dp.runs: dp.runs[0].italic = True
# #                 doc.paragraphs[idx]._element.addnext(dp._element)
# #                 idx += 1

# #             # Description
# #             if job.get("description"):
# #                 dp2 = doc.add_paragraph(job["description"])
# #                 doc.paragraphs[idx]._element.addnext(dp2._element)
# #                 idx += 1

# #             # Achievements (ensure we treat these as bullets)
# #             achievements = job.get("achievements", [])
# #             if isinstance(achievements, list):
# #                 for ach in achievements:
# #                     if ach.strip():
# #                         # Explicitly use List Bullet style
# #                         a = doc.add_paragraph(ach, style="List Bullet")
# #                         doc.paragraphs[idx]._element.addnext(a._element)
# #                         idx += 1
            
# #             # Spacer
# #             spacer = doc.add_paragraph("")
# #             doc.paragraphs[idx]._element.addnext(spacer._element)
# #             idx += 1

# #     # Education
# #     edu_lines = []
# #     for edu in data.get("education", []):
# #         parts = []
# #         if edu.get("degree"): parts.append(edu["degree"])
# #         if edu.get("field"): parts.append(f"in {edu['field']}")
# #         if edu.get("institution"): parts.append(f"| {edu['institution']}")
# #         if edu.get("end_year"): parts.append(f"| {edu['end_year']}")
        
# #         line = " ".join(parts)
# #         if line:
# #             edu_lines.append(line)
    
# #     if edu_lines:
# #         add_after_heading("Education", edu_lines, bullet=True)

# #     # --- 3. Handle Certifications ---
# #     certifications = data.get("certifications", [])
# #     cert_lines = [c.strip() for c in certifications if c and c.strip()]

# #     # Clean up empty cert blocks
# #     indices_to_remove = []
# #     for i, p in enumerate(doc.paragraphs):
# #         if "certification" in p.text.strip().lower():
# #             if not cert_lines:
# #                 indices_to_remove.append(i)
# #                 if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i+1].text.strip():
# #                     indices_to_remove.append(i+1)
# #             break 
    
# #     for idx in sorted(indices_to_remove, reverse=True):
# #         p = doc.paragraphs[idx]._element
# #         p.getparent().remove(p)

# #     if cert_lines:
# #         add_after_heading("CERTIFICATION", cert_lines, bullet=True)
    
# #     # --- 4. Save DOCX ---
# #     os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
# #     doc.save(output_docx_path)
# #     print(f"✅ DOCX saved: {output_docx_path}")

# #     # --- 5. Generate PDF ---
# #     create_pdf_from_docx(output_docx_path, output_pdf_path)


# # def find_dejavu_ttf():
# #     """Return a path to a DejaVuSans.ttf if present on the system or None."""
# #     candidates = [
# #         "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
# #         "/Library/Fonts/DejaVuSans.ttf",
# #         "C:\\Windows\\Fonts\\DejaVuSans.ttf",
# #         os.path.join(os.getcwd(), "DejaVuSans.ttf"),
# #     ]
# #     for p in candidates:
# #         if p and os.path.exists(p):
# #             return p
# #     return None


# # def create_pdf_from_docx(docx_path: str, output_pdf_path: str):
# #     """
# #     Create a PDF from a DOCX file with improved style detection logic.
# #     """
    
# #     # 1. Try docx2pdf (Best Fidelity)
# #     if _HAS_DOCX2PDF and platform.system() in ("Windows", "Darwin"):
# #         try:
# #             print("ℹ Using docx2pdf for conversion (best fidelity)")
# #             os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
# #             docx2pdf_convert(docx_path, output_pdf_path)
# #             return
# #         except Exception as e:
# #             print(f"⚠ docx2pdf conversion failed: {e} — falling back to text-render PDF")

# #     # 2. Fallback: FPDF with Formatting Logic
# #     print("ℹ Falling back to FPDF with manual style replication")
    
# #     try:
# #         doc = Document(docx_path)
# #     except Exception as e:
# #         print(f"❌ Failed to read temporary DOCX for PDF generation: {e}")
# #         return

# #     pdf = FPDF(format='A4', unit='mm')
# #     pdf.set_auto_page_break(auto=True, margin=15)
# #     pdf.add_page()

# #     # Font Setup
# #     dejavu = find_dejavu_ttf()
# #     use_unicode = False
# #     font_family = "Arial"
    
# #     if dejavu:
# #         try:
# #             # Register generic DejaVu
# #             pdf.add_font("DejaVu", "", dejavu, uni=True)
# #             # Register fake Bold version pointing to same file (works in recent FPDF)
# #             pdf.add_font("DejaVu", "B", dejavu, uni=True) 
# #             font_family = "DejaVu"
# #             use_unicode = True
# #             print(f"ℹ Registered DejaVu font at: {dejavu}")
# #         except:
# #             # Fallback for old FPDF
# #             try:
# #                 pdf.add_font("DejaVu", "", dejavu)
# #                 font_family = "DejaVu"
# #                 use_unicode = True
# #             except:
# #                 pass

# #     # Effective Width
# #     try:
# #         effective_width = pdf.epw
# #     except:
# #         effective_width = 190

# #     # --- Render Paragraphs with Style Detection ---
# #     for para in doc.paragraphs:
# #         text = para.text.strip()
# #         if not text:
# #             pdf.ln(5) # Add spacing for empty paragraphs
# #             continue

# #         # Defaults
# #         style = ''
# #         size = 10
# #         align = 'L'
# #         prefix = ''

# #         # 1. Detect Bullet Points (List Bullet Style)
# #         if 'list' in para.style.name.lower() or 'bullet' in para.style.name.lower():
# #             prefix = "- " # Simulate bullet
        
# #         # 2. Detect Headings / Titles
# #         if 'title' in para.style.name.lower():
# #             style = 'B'
# #             size = 14
# #             align = 'C'
# #         elif 'heading' in para.style.name.lower():
# #             style = 'B'
# #             size = 11
# #             # Add extra space before headings
# #             pdf.ln(2)

# #         # 3. Detect Centered Text (Alignment)
# #         if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
# #             align = 'C'

# #         # 4. Detect Manual Bold (e.g. "Company Name" lines)
# #         # Check if the first run in the paragraph is bold
# #         if not style and para.runs and para.runs[0].bold:
# #             style = 'B'

# #         # 5. Detect Headers that are just UPPERCASE lines (common in resumes)
# #         if not style and text.isupper() and len(text) < 50:
# #             style = 'B'
# #             size = 11
# #             pdf.ln(2) # Spacing before section headers

# #         # Apply settings
# #         pdf.set_font(font_family, style, size)
        
# #         # Sanitize text
# #         final_text = prefix + text
# #         if not use_unicode:
# #             final_text = clean_text(final_text)

# #         try:
# #             pdf.multi_cell(w=effective_width, h=5, txt=final_text, align=align)
# #         except Exception as e:
# #             # Emergency ASCII fallback
# #             safe = final_text.encode('ascii', 'ignore').decode('ascii')
# #             pdf.multi_cell(w=effective_width, h=5, txt=safe, align=align)

# #     os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
# #     pdf.output(output_pdf_path)
# #     if os.path.exists(output_pdf_path):
# #         print(f"✅ PDF saved: {output_pdf_path}")


import json
import os
import platform
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF helpers
from fpdf import FPDF

# Optional: docx2pdf (Windows/Mac) — used if available
try:
    from docx2pdf import convert as docx2pdf_convert
    _HAS_DOCX2PDF = True
except Exception:
    _HAS_DOCX2PDF = False


def clean_text(text):
    """
    Sanitizes text for FPDF (Latin-1) fallback.
    Replaces special Unicode characters that might cause crashes.
    """
    if not isinstance(text, str):
        return str(text) if text is not None else ""
    
    replacements = {
        '\u2013': '-',   # En-dash
        '\u2014': '--',  # Em-dash
        '\u2018': "'",   # Left single quote
        '\u2019': "'",   # Right single quote
        '\u201c': '"',   # Left double quote
        '\u201d': '"',   # Right double quote
        '\u2022': '*',   # Bullet
        '\u00a0': ' '    # Non-breaking space
    }
    
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    
    return text


def fill_template(template_path, output_docx_path, output_pdf_path, resume_json):
    """
    Fill the DOCX template (python-docx) and generate a readable PDF
    that mimics the DOCX formatting.
    """
    print("▶ fill_template start")

    # --- 1. Parse Data Robustly ---
    if isinstance(resume_json, str):
        try:
            data = json.loads(resume_json)
        except json.JSONDecodeError:
            print("❌ Error: Invalid JSON string provided.")
            return
    else:
        data = resume_json if resume_json else {}

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"❌ Error loading template: {e}")
        return

    # --- Helper Inner Functions ---
    def find_heading_idx(heading_text):
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip().lower() == heading_text.lower():
                return idx
        return None

    def clear_paragraph(para: Paragraph):
        for run in para.runs:
            run.text = ""
        para.text = ""

    def add_after_heading(heading, lines, bullet=False):
        idx = find_heading_idx(heading)
        if idx is not None:
            # Skip placeholder empty paragraphs immediately after header
            while idx + 1 < len(doc.paragraphs) and not doc.paragraphs[idx + 1].text.strip():
                # If we find an empty placeholder, we can clear it or skip it.
                # To be safe and clean, let's clear it so it doesn't take up whitespace.
                clear_paragraph(doc.paragraphs[idx + 1])
                idx += 1
            
            # Insert new paragraphs
            for line in lines:
                if not line.strip(): continue
                # We apply "List Bullet" style here so the PDF generator can detect it later
                style = "List Bullet" if bullet else None
                new_para = doc.add_paragraph(line, style=style)
                # Insert paragraph after the current index
                doc.paragraphs[idx]._element.addnext(new_para._element)
                idx += 1

    # --- 2. Fill Core Fields ---

    # Full Name
    for para in doc.paragraphs:
        if "full name" in para.text.lower():
            full_name = f"{data.get('first_name', '')} {data.get('last_name', '')}".strip()
            clear_paragraph(para)
            para.add_run(full_name)
            # Explicitly center and bold the name for DOCX
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if para.runs: para.runs[0].bold = True
            break

    # --- UPDATED: Career Summary (Now Bullets) ---
    summary_text = data.get("career_summary", "")
    summary_lines = []
    if summary_text:
        # 1. Check if it already has newlines
        if '\n' in summary_text:
            raw_lines = summary_text.split('\n')
        else:
            # 2. If no newlines, split by sentences (. ) to make bullets
            raw_lines = summary_text.split('. ')
            
        for line in raw_lines:
            clean_line = line.strip()
            if clean_line:
                # Ensure punctuation remains if we split by it
                if not clean_line.endswith('.') and len(clean_line) > 5:
                    clean_line += '.'
                summary_lines.append(clean_line)

    if summary_lines:
        # Use helper with bullet=True
        add_after_heading("Career Summary", summary_lines, bullet=True)


    # Expertise (Joined by Pipe |)
    idx = find_heading_idx("Expertise")
    if idx is not None and idx + 1 < len(doc.paragraphs):
        clear_paragraph(doc.paragraphs[idx + 1])
        expertise_list = data.get("expertise", [])
        if expertise_list:
            doc.paragraphs[idx + 1].add_run(" | ".join(expertise_list))

    # Technical Skills (Joined by Pipe |)
    idx = find_heading_idx("Technical Skills")
    if idx is not None and idx + 1 < len(doc.paragraphs):
        clear_paragraph(doc.paragraphs[idx + 1])
        skills_list = data.get("technical_skills", [])
        if skills_list:
            doc.paragraphs[idx + 1].add_run(" | ".join(skills_list))

    # Professional Experience
    idx = find_heading_idx("Professional Experience")
    if idx is not None:
        for job in data.get("professional_experience", []):
            if not isinstance(job, dict): continue

            # Header: Title | Company | Location
            header_parts = [job.get("title", ""), job.get("company", ""), job.get("location", "")]
            header_parts = [p for p in header_parts if p]
            
            if header_parts:
                hp = doc.add_paragraph(" | ".join(header_parts))
                if hp.runs: hp.runs[0].bold = True
                doc.paragraphs[idx]._element.addnext(hp._element)
                idx += 1

            # Dates
            date_parts = [job.get("start_date", ""), job.get("end_date", "")]
            date_parts = [d for d in date_parts if d]
            if date_parts:
                dp = doc.add_paragraph(" – ".join(date_parts)) 
                if dp.runs: dp.runs[0].italic = True
                doc.paragraphs[idx]._element.addnext(dp._element)
                idx += 1

            # Description
            if job.get("description"):
                dp2 = doc.add_paragraph(job["description"])
                doc.paragraphs[idx]._element.addnext(dp2._element)
                idx += 1

            # Achievements (ensure we treat these as bullets)
            achievements = job.get("achievements", [])
            if isinstance(achievements, list):
                for ach in achievements:
                    if ach.strip():
                        # Explicitly use List Bullet style
                        a = doc.add_paragraph(ach, style="List Bullet")
                        doc.paragraphs[idx]._element.addnext(a._element)
                        idx += 1
            
            # Spacer
            spacer = doc.add_paragraph("")
            doc.paragraphs[idx]._element.addnext(spacer._element)
            idx += 1

    # Education
    edu_lines = []
    for edu in data.get("education", []):
        parts = []
        if edu.get("degree"): parts.append(edu["degree"])
        if edu.get("field"): parts.append(f"in {edu['field']}")
        if edu.get("institution"): parts.append(f"| {edu['institution']}")
        if edu.get("end_year"): parts.append(f"| {edu['end_year']}")
        
        line = " ".join(parts)
        if line:
            edu_lines.append(line)
    
    if edu_lines:
        add_after_heading("Education", edu_lines, bullet=True)

    # --- 3. Handle Certifications ---
    certifications = data.get("certifications", [])
    cert_lines = [c.strip() for c in certifications if c and c.strip()]

    # Clean up empty cert blocks
    indices_to_remove = []
    for i, p in enumerate(doc.paragraphs):
        if "certification" in p.text.strip().lower():
            if not cert_lines:
                indices_to_remove.append(i)
                if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i+1].text.strip():
                    indices_to_remove.append(i+1)
            break 
    
    for idx in sorted(indices_to_remove, reverse=True):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)

    if cert_lines:
        add_after_heading("CERTIFICATION", cert_lines, bullet=True)
    
    # --- 4. Save DOCX ---
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    doc.save(output_docx_path)
    print(f"✅ DOCX saved: {output_docx_path}")

    # --- 5. Generate PDF ---
    create_pdf_from_docx(output_docx_path, output_pdf_path)


def find_dejavu_ttf():
    """Return a path to a DejaVuSans.ttf if present on the system or None."""
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/DejaVuSans.ttf",
        "C:\\Windows\\Fonts\\DejaVuSans.ttf",
        os.path.join(os.getcwd(), "DejaVuSans.ttf"),
    ]
    for p in candidates:
        if p and os.path.exists(p):
            return p
    return None


def create_pdf_from_docx(docx_path: str, output_pdf_path: str):
    """
    Create a PDF from a DOCX file with improved style detection logic.
    """
    
    # 1. Try docx2pdf (Best Fidelity)
    if _HAS_DOCX2PDF and platform.system() in ("Windows", "Darwin"):
        try:
            print("ℹ Using docx2pdf for conversion (best fidelity)")
            os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
            docx2pdf_convert(docx_path, output_pdf_path)
            return
        except Exception as e:
            print(f"⚠ docx2pdf conversion failed: {e} — falling back to text-render PDF")

    # 2. Fallback: FPDF with Formatting Logic
    print("ℹ Falling back to FPDF with manual style replication")
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"❌ Failed to read temporary DOCX for PDF generation: {e}")
        return

    pdf = FPDF(format='A4', unit='mm')
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Font Setup
    dejavu = find_dejavu_ttf()
    use_unicode = False
    font_family = "Arial"
    
    if dejavu:
        try:
            # Register generic DejaVu
            pdf.add_font("DejaVu", "", dejavu, uni=True)
            # Register fake Bold version pointing to same file (works in recent FPDF)
            pdf.add_font("DejaVu", "B", dejavu, uni=True) 
            font_family = "DejaVu"
            use_unicode = True
            print(f"ℹ Registered DejaVu font at: {dejavu}")
        except:
            # Fallback for old FPDF
            try:
                pdf.add_font("DejaVu", "", dejavu)
                font_family = "DejaVu"
                use_unicode = True
            except:
                pass

    # Effective Width
    try:
        effective_width = pdf.epw
    except:
        effective_width = 190

    # --- Render Paragraphs with Style Detection ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            pdf.ln(5) # Add spacing for empty paragraphs
            continue

        # Defaults
        style = ''
        size = 10
        align = 'L'
        prefix = ''

        # 1. Detect Bullet Points (List Bullet Style)
        if 'list' in para.style.name.lower() or 'bullet' in para.style.name.lower():
            prefix = "- " # Simulate bullet
        
        # 2. Detect Headings / Titles
        if 'title' in para.style.name.lower():
            style = 'B'
            size = 14
            align = 'C'
        elif 'heading' in para.style.name.lower():
            style = 'B'
            size = 11
            # Add extra space before headings
            pdf.ln(2)

        # 3. Detect Centered Text (Alignment)
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align = 'C'

        # 4. Detect Manual Bold (e.g. "Company Name" lines)
        # Check if the first run in the paragraph is bold
        if not style and para.runs and para.runs[0].bold:
            style = 'B'

        # 5. Detect Headers that are just UPPERCASE lines (common in resumes)
        if not style and text.isupper() and len(text) < 50:
            style = 'B'
            size = 11
            pdf.ln(2) # Spacing before section headers

        # Apply settings
        pdf.set_font(font_family, style, size)
        
        # Sanitize text
        final_text = prefix + text
        if not use_unicode:
            final_text = clean_text(final_text)

        try:
            pdf.multi_cell(w=effective_width, h=5, txt=final_text, align=align)
        except Exception as e:
            # Emergency ASCII fallback
            safe = final_text.encode('ascii', 'ignore').decode('ascii')
            pdf.multi_cell(w=effective_width, h=5, txt=safe, align=align)

    os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
    pdf.output(output_pdf_path)
    if os.path.exists(output_pdf_path):
        print(f"✅ PDF saved: {output_pdf_path}")